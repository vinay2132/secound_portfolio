"""
Document processing utilities for PDF, DOCX, and TXT files
"""

import streamlit as st
import PyPDF2
import docx
from pathlib import Path
from config.constants import DEFAULT_RESUME, DEFAULT_PROJECTS


def extract_text_from_pdf(file_path):
    """Extract text from PDF file"""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
            return text
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return None


def extract_text_from_pdf_upload(file):
    """Extract text from uploaded PDF file"""
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text


def extract_text_from_docx(file):
    """Extract text from DOCX file"""
    doc = docx.Document(file)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return text


def extract_text_from_txt(file):
    """Extract text from TXT file"""
    return file.read().decode('utf-8')


def process_uploaded_file(file):
    """Process uploaded file and extract text"""
    file_extension = file.name.split('.')[-1].lower()
    
    if file_extension == 'pdf':
        return extract_text_from_pdf_upload(file)
    elif file_extension == 'docx':
        return extract_text_from_docx(file)
    elif file_extension == 'txt':
        return extract_text_from_txt(file)
    else:
        return None


def auto_load_default_resume():
    """Automatically load the default resume from the current directory"""
    current_dir = Path('.')
    resume_path = current_dir / DEFAULT_RESUME
    
    if resume_path.exists() and DEFAULT_RESUME not in st.session_state.documents:
        text = extract_text_from_pdf(resume_path)
        if text:
            st.session_state.documents[DEFAULT_RESUME] = text
            return True
    elif DEFAULT_RESUME in st.session_state.documents:
        return True
    
    return False


def auto_load_projects_file():
    """Automatically load the projects file from the current directory"""
    current_dir = Path('.')
    projects_path = current_dir / DEFAULT_PROJECTS
    
    if projects_path.exists() and DEFAULT_PROJECTS not in st.session_state.documents:
        try:
            text = projects_path.read_text(encoding='utf-8')
            if text:
                st.session_state.documents[DEFAULT_PROJECTS] = text
                return True
        except Exception as e:
            st.error(f"Error reading projects file: {str(e)}")
    elif DEFAULT_PROJECTS in st.session_state.documents:
        return True
    
    return False


def auto_load_additional_documents():
    """Automatically load any additional documents from the current directory"""
    current_dir = Path('.')
    loaded_count = 0
    
    for ext in ['*.pdf', '*.docx', '*.txt']:
        for file_path in current_dir.glob(ext):
            if file_path.name not in st.session_state.documents and file_path.name not in [DEFAULT_RESUME, DEFAULT_PROJECTS]:
                try:
                    if ext == '*.pdf':
                        text = extract_text_from_pdf(file_path)
                    elif ext == '*.docx':
                        doc = docx.Document(file_path)
                        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    else:
                        text = file_path.read_text(encoding='utf-8')
                    
                    if text:
                        st.session_state.documents[file_path.name] = text
                        loaded_count += 1
                except Exception as e:
                    continue
    
    return loaded_count

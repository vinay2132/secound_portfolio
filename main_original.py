import streamlit as st
import google.generativeai as genai
from pathlib import Path
import PyPDF2
import docx
import json
from datetime import datetime
import os
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Page configuration
st.set_page_config(
    page_title="AI Career Assistant",
    page_icon="💼",
    layout="wide"
)

# Initialize session state
if 'documents' not in st.session_state:
    st.session_state.documents = {}
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'auto_loaded' not in st.session_state:
    st.session_state.auto_loaded = False
if 'job_description' not in st.session_state:
    st.session_state.job_description = ""
if 'jd_configured' not in st.session_state:
    st.session_state.jd_configured = False
if 'personal_details' not in st.session_state:
    st.session_state.personal_details = """
Full Name: Vinay Ramesh
Email: vinayramesh6020@gmail.com
Contact Number: +1 (682) 273-5833
Education: Master of Science in Data Science (UNT, Sep 2023 – Dec 2024), Bachelor of Technology in Computer Science (TKR College of Engineering & Technology, Jun 2017 – May 2021)
Current Location: Denton, TX
Relocation: Yes
Availability to Take Call: Yes
Availability to Join: Within 1-2 weeks
Relevant Experience: 3+ years in Full Stack Development
Visa/Work Authorization: F1 OPT
"""
if 'writing_guidelines' not in st.session_state:
    st.session_state.writing_guidelines = """
CRITICAL WRITING GUIDELINES - MUST FOLLOW FOR ALL OUTPUTS:

1. TONE & STYLE:
   - Write as a seasoned professional with 3+ years of experience
   - Make it sound HUMAN and natural, NOT AI-generated
   - Be professional, confident, and recruiter-friendly
   - Avoid over-formal or robotic language

2. FORMATTING RULES:
   - DO NOT use highlighting, bold text, asterisks (**), or excessive markdown
   - Keep content clean and simple
   - For emails: Include a relevant subject line at the top
   - Use proper paragraph spacing, but keep it minimal

3. EMAIL-SPECIFIC RULES:
   - Keep emails SHORT (1-2 concise paragraphs maximum)
   - Highlight specific technologies from BOTH job description and resume
   - If no hiring manager name is provided, use: "Dear Hiring Manager,"
   - ALWAYS mention F1 OPT work authorization clearly when relevant
   - ALWAYS use this exact signature format:

--
Best regards,
Vinay Ramesh
Senior Full Stack Developer
📧 vinayramesh6020@gmail.com
📞 +1 (682) 273-5833

4. CONTENT REQUIREMENTS:
   - Incorporate personal details naturally where relevant
   - Match technologies and skills to the job requirements
   - Be specific and avoid generic phrases
   - Show genuine interest and understanding of the role

5. GENERAL RULES FOR ALL OUTPUTS:
   - Prioritize clarity and readability
   - Use active voice
   - Be concise and to the point
   - Proofread for grammar and professionalism
"""

# Default resume filename
DEFAULT_RESUME = "Vinay_Ramesh_full_stack_developer.pdf"

# Helper Functions
def extract_text_from_pdf(file_path):
    """Extract text from PDF file"""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
            return text
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return None

def extract_text_from_pdf_upload(file):
    """Extract text from uploaded PDF file"""
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_docx(file):
    """Extract text from DOCX file"""
    doc = docx.Document(file)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return text

def extract_text_from_txt(file):
    """Extract text from TXT file"""
    return file.read().decode('utf-8')

def process_uploaded_file(file):
    """Process uploaded file and extract text"""
    file_extension = file.name.split('.')[-1].lower()
    
    if file_extension == 'pdf':
        return extract_text_from_pdf_upload(file)
    elif file_extension == 'docx':
        return extract_text_from_docx(file)
    elif file_extension == 'txt':
        return extract_text_from_txt(file)
    else:
        return None

def auto_load_default_resume():
    """Automatically load the default resume from the current directory"""
    current_dir = Path('.')
    resume_path = current_dir / DEFAULT_RESUME
    
    if resume_path.exists() and DEFAULT_RESUME not in st.session_state.documents:
        text = extract_text_from_pdf(resume_path)
        if text:
            st.session_state.documents[DEFAULT_RESUME] = text
            return True
    elif DEFAULT_RESUME in st.session_state.documents:
        return True
    
    return False

def auto_load_additional_documents():
    """Automatically load any additional documents from the current directory"""
    current_dir = Path('.')
    loaded_count = 0
    
    for ext in ['*.pdf', '*.docx', '*.txt']:
        for file_path in current_dir.glob(ext):
            if file_path.name not in st.session_state.documents and file_path.name != DEFAULT_RESUME:
                try:
                    if ext == '*.pdf':
                        text = extract_text_from_pdf(file_path)
                    elif ext == '*.docx':
                        doc = docx.Document(file_path)
                        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    else:
                        text = file_path.read_text(encoding='utf-8')
                    
                    if text:
                        st.session_state.documents[file_path.name] = text
                        loaded_count += 1
                except Exception as e:
                    continue
    
    return loaded_count

def get_context():
    """Build context from all uploaded documents, personal details, job description, and guidelines"""
    context = f"""{st.session_state.writing_guidelines}

PERSONAL DETAILS:
{st.session_state.personal_details}

TARGET JOB DESCRIPTION:
{st.session_state.job_description}

RESUME AND DOCUMENTS:
"""
    for doc_name, doc_content in st.session_state.documents.items():
        context += f"\n--- {doc_name} ---\n{doc_content}\n"
    
    return context

def call_gemini(prompt, api_key):
    """Call Gemini API with the given prompt"""
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(model_name='gemini-2.0-flash-exp')
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Error: {str(e)}"

# Auto-load documents on first run
if not st.session_state.auto_loaded:
    resume_loaded = auto_load_default_resume()
    additional_loaded = auto_load_additional_documents()
    st.session_state.auto_loaded = True

# Get API key from environment or user input
env_api_key = os.getenv('GEMINI_API_KEY')
DEFAULT_API_KEY = "AIzaSyD26rG9VpYFPmYCnQ_U8OO6jOKREzEZRYY"

# Sidebar for API Key and Document Upload
with st.sidebar:
    st.title("⚙️ Configuration")
    
    # API Key Input
    if env_api_key:
        st.success("✅ API Key loaded from .env file")
        api_key = env_api_key
        masked_key = env_api_key[:8] + "..." + env_api_key[-4:]
        st.text(f"Key: {masked_key}")
    else:
        api_key = DEFAULT_API_KEY
        st.success("✅ Using default API Key")
        masked_key = DEFAULT_API_KEY[:8] + "..." + DEFAULT_API_KEY[-4:]
        st.text(f"Key: {masked_key}")
        
        custom_key = st.text_input("Override API Key (optional)", type="password", 
                                 help="Leave empty to use default key")
        if custom_key:
            api_key = custom_key
    
    st.divider()
    
    # Job Description Configuration
    with st.expander("🎯 Configure Target Job Description", expanded=not st.session_state.jd_configured):
        st.markdown("**Enter the job description once - it will be used across all features:**")
        job_desc_input = st.text_area(
            "Job Description",
            value=st.session_state.job_description,
            height=300,
            placeholder="""Paste the complete job description here including:
- Job title and company name
- Requirements and qualifications
- Technologies and skills needed
- Job responsibilities
- Any other relevant details

This will be used for all email generation, resume updates, and cover letters.""",
            help="This job description will be the default context for all operations"
        )
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("💾 Save Job Description", use_container_width=True):
                st.session_state.job_description = job_desc_input
                st.session_state.jd_configured = True
                st.success("Job description saved! All features will now use this as context.")
                st.rerun()
        with col2:
            if st.button("🗑️ Clear Job Description", use_container_width=True):
                st.session_state.job_description = ""
                st.session_state.jd_configured = False
                st.info("Job description cleared.")
                st.rerun()
        
        if st.session_state.jd_configured:
            st.success("✅ Job description is configured and active!")
    
    st.divider()
    
    # Personal Details Editor
    with st.expander("✏️ Edit Personal Details"):
        st.markdown("**Update your personal information:**")
        personal_details_input = st.text_area(
            "Personal Details",
            value=st.session_state.personal_details,
            height=300,
            help="Edit your personal details that will be used in all generated content"
        )
        if st.button("💾 Save Personal Details"):
            st.session_state.personal_details = personal_details_input
            st.success("Personal details updated!")
            st.rerun()
    
    st.divider()
    
    # Show loaded documents status
    if st.session_state.documents:
        if DEFAULT_RESUME in st.session_state.documents:
            st.success(f"✅ Default resume loaded: {DEFAULT_RESUME}")
        
        additional_docs = len(st.session_state.documents) - (1 if DEFAULT_RESUME in st.session_state.documents else 0)
        if additional_docs > 0:
            st.info(f"📄 {additional_docs} additional document(s) loaded")
    else:
        st.warning(f"⚠️ Default resume not found: {DEFAULT_RESUME}")
        st.info("Please upload your resume manually below.")
    
    # Document Upload Section
    st.subheader("📁 Upload Documents")
    uploaded_files = st.file_uploader(
        "Upload Resume or Additional Documents",
        type=['pdf', 'docx', 'txt'],
        accept_multiple_files=True,
        help="Upload your resume or additional documents"
    )
    
    if uploaded_files:
        for file in uploaded_files:
            if file.name not in st.session_state.documents:
                text = process_uploaded_file(file)
                if text:
                    st.session_state.documents[file.name] = text
                    st.success(f"✅ {file.name} uploaded!")
    
    # Display all loaded documents
    if st.session_state.documents:
        st.divider()
        st.subheader("📚 Loaded Documents")
        for doc_name in st.session_state.documents.keys():
            col1, col2 = st.columns([3, 1])
            with col1:
                if doc_name == DEFAULT_RESUME:
                    st.text(f"⭐ {doc_name}")
                else:
                    st.text(doc_name)
            with col2:
                if st.button("🗑️", key=f"delete_{doc_name}"):
                    del st.session_state.documents[doc_name]
                    st.rerun()

# Main Content
st.title("💼 AI Career Assistant")
st.markdown("*Your personal AI assistant for career management powered by Gemini*")

# Check for API key
if not api_key:
    st.warning("⚠️ Please enter your Gemini API key in the sidebar or add it to your .env file.")
    st.info("""
    ### How to setup:
    1. Create a `.env` file in the same folder
    2. Add this line: `GEMINI_API_KEY=your_api_key_here`
    3. Or enter the key manually in the sidebar
    
    Get your API key from [Google AI Studio](https://makersuite.google.com/app/apikey)
    """)
    st.stop()

# Check if documents are loaded
if not st.session_state.documents:
    st.error(f"⚠️ No documents loaded. Please make sure '{DEFAULT_RESUME}' is in the same folder as this script, or upload your resume manually in the sidebar.")
    st.info("""
    ### Getting Started:
    1. Place your resume file named `Vinay_Ramesh_full_stack_developer.pdf` in the same folder as this script
    2. Or upload your resume manually using the sidebar
    3. Configure your target job description in the sidebar
    4. Choose a task from the tabs below once your documents are loaded
    """)
    st.stop()

# Check if job description is configured
if not st.session_state.jd_configured or not st.session_state.job_description.strip():
    st.warning("⚠️ Please configure your target job description first!")
    st.info("""
    ### Before You Begin:
    
    **📋 Configure your target job description in the sidebar** (under "🎯 Configure Target Job Description")
    
    This is a **one-time setup** that will:
    - ✅ Generate tailored emails automatically
    - ✅ Update your resume to match the job requirements
    - ✅ Create customized cover letters
    - ✅ Provide job-specific career advice
    
    Simply paste the complete job posting once, and all features will use it as context!
    """)
    st.stop()

# Show job description status
st.success(f"✅ Target job configured! All features are using your job description as context.")
with st.expander("📋 View Current Job Description"):
    st.text(st.session_state.job_description[:500] + "..." if len(st.session_state.job_description) > 500 else st.session_state.job_description)

# Create tabs for different functionalities
tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "📧 Email Writer",
    "📄 Resume Updater",
    "✉️ Cover Letter",
    "💬 Q&A Assistant",
    "📊 Document Summary"
])

# Tab 1: Email Writer
with tab1:
    st.header("📧 Professional Email Writer")
    st.markdown("Generate professional emails based on your configured job description.")
    st.info("💡 Using your configured job description as context")
    
    col1, col2 = st.columns(2)
    with col1:
        email_purpose = st.selectbox(
            "Email Purpose",
            ["Job Application", "Follow-up", "Networking", "Thank You", "Custom"]
        )
    with col2:
        email_tone = st.selectbox(
            "Tone",
            ["Professional", "Confident", "Friendly", "Formal"]
        )
    
    hiring_manager_name = st.text_input(
        "Hiring Manager Name (optional)",
        placeholder="Leave empty if unknown - will use 'Dear Hiring Manager,'",
        help="Enter the hiring manager's name if known"
    )
    
    additional_context = st.text_area(
        "Additional Context (optional)",
        placeholder="Add any extra details or specific points you want to mention...",
        height=100
    )
    
    if st.button("✨ Generate Email", key="generate_email"):
        with st.spinner("Generating email..."):
            user_context = get_context()
            
            salutation = f"Dear {hiring_manager_name}," if hiring_manager_name else "Dear Hiring Manager,"
            
            prompt = f"""
{user_context}

TASK: Write a {email_tone.lower()} job application email for: {email_purpose}

SALUTATION: {salutation}

ADDITIONAL CONTEXT (if provided):
{additional_context}

REQUIREMENTS:
- Start with: Subject: [create a short, relevant subject line]
- Use the salutation: {salutation}
- Keep it to 1-2 SHORT paragraphs maximum
- Highlight specific technologies that match BOTH the target job description and my resume
- Clearly mention F1 OPT work authorization
- End with the EXACT signature format from the guidelines
- Make it sound human, natural, and confident
- NO bold text, asterisks, or highlighting
- Reference the TARGET JOB DESCRIPTION provided in the context

Generate the email now:
"""
            
            email = call_gemini(prompt, api_key)
            st.markdown("### Generated Email:")
            st.text(email)
            
            # Download button
            st.download_button(
                label="📥 Download Email",
                data=email,
                file_name=f"email_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                mime="text/plain"
            )

# Tab 2: Resume Updater
with tab2:
    st.header("📄 Resume Updater")
    st.markdown("Update your resume to match your configured job description.")
    st.info("💡 Using your configured job description as context")
    
    update_type = st.radio(
        "What would you like to update?",
        ["Tailor entire resume", "Update skills section", "Optimize summary", "Highlight relevant projects"]
    )
    
    additional_instructions = st.text_area(
        "Additional Instructions (optional)",
        placeholder="Any specific points to emphasize or changes to make...",
        height=100
    )
    
    if st.button("🔄 Update Resume", key="update_resume"):
        with st.spinner("Updating resume..."):
            user_context = get_context()
            prompt = f"""
{user_context}

TASK: {update_type}

Based on my background and the TARGET JOB DESCRIPTION provided above, {update_type.lower()}.

ADDITIONAL INSTRUCTIONS:
{additional_instructions}

REQUIREMENTS:
- Highlight relevant skills and experiences that match the target job
- Use keywords from the job description naturally
- Maintain professional format WITHOUT excessive bold or highlighting
- Emphasize achievements that match the role
- Keep it concise and ATS-friendly
- Sound natural and human, not AI-generated

Provide the updated section or full resume:
"""
            
            updated_resume = call_gemini(prompt, api_key)
            st.markdown("### Updated Resume:")
            st.text(updated_resume)
            
            # Download button
            st.download_button(
                label="📥 Download Updated Resume",
                data=updated_resume,
                file_name=f"resume_updated_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                mime="text/plain"
            )

# Tab 3: Cover Letter Generator
with tab3:
    st.header("✉️ Cover Letter Generator")
    st.markdown("Create compelling cover letters based on your configured job description.")
    st.info("💡 Using your configured job description as context")
    
    hiring_manager_cl = st.text_input(
        "Hiring Manager Name (optional)",
        placeholder="Leave empty if unknown",
        help="Enter the hiring manager's name if known",
        key="cl_hiring_manager"
    )
    
    why_interested = st.text_area(
        "Why are you interested in this position? (optional)",
        placeholder="Share your specific interest in this role or company...",
        height=100
    )
    
    if st.button("✨ Generate Cover Letter", key="generate_cl"):
        with st.spinner("Crafting your cover letter..."):
            user_context = get_context()
            
            prompt = f"""
{user_context}

TASK: Write a compelling cover letter for the TARGET JOB DESCRIPTION provided above

HIRING MANAGER: {hiring_manager_cl if hiring_manager_cl else "Not provided - use 'Dear Hiring Manager,'"}

WHY INTERESTED (if provided):
{why_interested}

INSTRUCTIONS:
Write a cover letter that:
- Uses "Dear Hiring Manager," if no specific name is provided, otherwise use the hiring manager's name
- Shows genuine interest in the role and company
- Highlights relevant experiences from my background that match the job requirements
- Matches specific technologies and skills from my resume to the job requirements
- Is professional yet personable and human-sounding
- Mentions F1 OPT work authorization clearly when relevant
- Includes proper formatting: date, salutation, body paragraphs, closing
- Avoids generic phrases - make it authentic and specific
- Uses NO excessive bold text or highlighting
- Ends with my complete contact signature from the guidelines
- Keeps it concise but compelling (2-3 paragraphs maximum)

Write the cover letter now:
"""
            
            cover_letter = call_gemini(prompt, api_key)
            st.markdown("### Your Cover Letter:")
            st.text(cover_letter)
            
            # Download button
            st.download_button(
                label="📥 Download Cover Letter",
                data=cover_letter,
                file_name=f"cover_letter_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                mime="text/plain"
            )

# Tab 4: Q&A Assistant
with tab4:
    st.header("💬 Career Q&A Assistant")
    st.markdown("Ask questions about the job description, your resume, or career advice.")
    st.info("💡 All responses will consider your configured job description")
    
    # Display chat history
    for message in st.session_state.chat_history:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
    
    # Chat input
    if question := st.chat_input("Ask me anything about your career, the target job, resume match, etc..."):
        # Add user message to chat
        st.session_state.chat_history.append({"role": "user", "content": question})
        with st.chat_message("user"):
            st.markdown(question)
        
        # Generate response
        with st.chat_message("assistant"):
            with st.spinner("Thinking..."):
                user_context = get_context()
                prompt = f"""
{user_context}

USER QUESTION: {question}

Based on my background, documents, personal details, and the TARGET JOB DESCRIPTION, provide a helpful, accurate, and personalized answer.
- If the question is about the target job, provide specific insights about how I match the requirements
- If comparing my profile to the job description, be specific about strengths and areas to highlight
- Keep responses natural, conversational, and professional
- Mention F1 OPT if work authorization is relevant to the question

Answer the question:
"""
                
                answer = call_gemini(prompt, api_key)
                st.markdown(answer)
                
                # Add assistant response to chat
                st.session_state.chat_history.append({"role": "assistant", "content": answer})
    
    # Clear chat button
    if st.button("🗑️ Clear Chat History"):
        st.session_state.chat_history = []
        st.rerun()

# Tab 5: Document Summary
with tab5:
    st.header("📊 Document Summary & Analysis")
    st.markdown("Get insights and summaries based on your documents and target job.")
    st.info("💡 Analysis will include job match assessment")
    
    analysis_type = st.selectbox(
        "What would you like to analyze?",
        [
            "Job match analysis",
            "Summarize all documents",
            "Extract key skills matching the job",
            "List relevant projects for this job",
            "Identify strengths and gaps for this role",
            "Generate career summary for this position"
        ]
    )
    
    if st.button("📊 Analyze", key="analyze_docs"):
        with st.spinner("Analyzing your documents..."):
            user_context = get_context()
            prompt = f"""
{user_context}

TASK: {analysis_type}

Provide a comprehensive analysis based on the request, considering:
- My resume and documents
- The TARGET JOB DESCRIPTION
- How well I match the job requirements

Be specific, detailed, and actionable:
- Keep formatting clean and minimal
- Make it professional and easy to read
- Avoid excessive bold text or highlighting
- Provide concrete recommendations where applicable

Generate the analysis:
"""
            
            analysis = call_gemini(prompt, api_key)
            st.markdown("### Analysis Results:")
            st.markdown(analysis)
            
            # Download button
            st.download_button(
                label="📥 Download Analysis",
                data=analysis,
                file_name=f"analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                mime="text/plain"
            )

# Footer
st.divider()
st.markdown("""
<div style='text-align: center; color: #666; padding: 20px;'>
    <p>💼 AI Career Assistant | Powered by Google Gemini</p>
    <p style='font-size: 0.8em;'>Your data is processed securely and not stored permanently.</p>
</div>
""", unsafe_allow_html=True)
"""
Advanced Resume Formatter - 3-Stage DOCX Generation
Integrates the multi-stage resume tailoring system into the Streamlit app
"""

import streamlit as st
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.table import CT_TblPr
from utils.gemini_api import call_gemini


def render_advanced_resume_formatter(api_key):
    """Render the Advanced Resume Formatter tab with 3-stage processing"""
    
    st.header("🎯 Advanced Resume Formatter (2-Page DOCX)")
    st.markdown("""
    **Generate a professionally formatted 2-page DOCX resume** tailored to your job description.
    
    This advanced system uses a **3-stage process**:
    1. ✨ **Stage 1**: Tailors content to match job requirements
    2. 🎨 **Stage 2**: Matches your natural writing tone and style  
    3. 📄 **Stage 3**: Optimizes for clean 2-page formatting
    """)
    
    st.info("💡 Using your configured job description and resume documents")
    
    # Check if documents are loaded
    if not st.session_state.documents:
        st.warning("⚠️ No documents loaded. Please upload your resume first.")
        return
    
    # Options
    with st.expander("⚙️ Formatting Options", expanded=False):
        col1, col2 = st.columns(2)
        
        with col1:
            include_summary = st.checkbox("Include Summary Section", value=True)
            include_skills = st.checkbox("Include Skills Section", value=True)
            
        with col2:
            include_experience = st.checkbox("Include Experience Section", value=True)
            include_education = st.checkbox("Include Education Section", value=True)
        
        st.markdown("**Contact Information:**")
        col3, col4 = st.columns(2)
        with col3:
            contact_name = st.text_input("Name", value="Vinay Ramesh")
            contact_phone = st.text_input("Phone", value="+1 (682) 273-5833")
        with col4:
            contact_email = st.text_input("Email", value="vinayramesh6020@gmail.com")
            max_bullets_per_job = st.number_input("Max bullets per job", min_value=3, max_value=8, value=5)
    
    # Additional instructions
    additional_notes = st.text_area(
        "Additional Instructions (optional)",
        placeholder="Any specific requirements or points to emphasize...",
        height=100
    )
    
    # Generate button
    if st.button("🚀 Generate 2-Page DOCX Resume", type="primary", use_container_width=True):
        
        # Get resume content from documents
        from config.constants import DEFAULT_RESUME
        resume_text = st.session_state.documents.get(DEFAULT_RESUME, "")
        
        if not resume_text:
            # Try to find any PDF document
            for doc_name, content in st.session_state.documents.items():
                if doc_name.endswith('.pdf'):
                    resume_text = content
                    break
        
        if not resume_text:
            st.error("❌ No resume document found. Please upload your resume.")
            return
        
        job_desc = st.session_state.job_description
        
        if not job_desc:
            st.error("❌ No job description configured. Please configure it in the sidebar.")
            return
        
        # Progress tracking
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        try:
            # Stage 1: Tailor content
            status_text.text("⏳ Stage 1/3: Tailoring content to job requirements...")
            progress_bar.progress(10)
            
            stage1_result = run_stage1_tailoring(resume_text, job_desc, api_key, additional_notes)
            progress_bar.progress(33)
            
            if not stage1_result:
                st.error("❌ Stage 1 failed. Please try again.")
                return
            
            # Save stage 1 for debugging
            st.session_state['stage1_output'] = stage1_result
            
            # Stage 2: Match tone
            status_text.text("⏳ Stage 2/3: Matching your natural writing style...")
            progress_bar.progress(40)
            
            stage2_result = run_stage2_tone_matching(resume_text, stage1_result, api_key)
            progress_bar.progress(66)
            
            if not stage2_result:
                st.error("❌ Stage 2 failed. Please try again.")
                return
            
            # Save stage 2 for debugging
            st.session_state['stage2_output'] = stage2_result
            
            # Stage 3: Optimize for 2 pages
            status_text.text("⏳ Stage 3/3: Optimizing for 2-page format...")
            progress_bar.progress(70)
            
            stage3_result = run_stage3_optimization(stage2_result, api_key, max_bullets_per_job)
            progress_bar.progress(90)
            
            if not stage3_result:
                st.error("❌ Stage 3 failed. Please try again.")
                return
            
            # Save stage 3 for debugging
            st.session_state['stage3_output'] = stage3_result
            
            # Parse and create DOCX
            status_text.text("⏳ Creating formatted DOCX file...")
            
            sections = parse_resume_content(stage3_result)
            
            # Apply section filters
            if not include_summary:
                sections['summary'] = ''
            if not include_skills:
                sections['skills'] = {}
            if not include_experience:
                sections['experience'] = []
            if not include_education:
                sections['education'] = ''
            
            # Create DOCX in memory
            doc = create_docx_resume_in_memory(
                sections, 
                contact_name, 
                contact_phone, 
                contact_email
            )
            
            progress_bar.progress(100)
            status_text.text("✅ Resume generation complete!")
            
            # Success message
            st.success("🎉 Your 2-page tailored resume is ready!")
            
            # Display stages in expanders
            col1, col2, col3 = st.columns(3)
            
            with col1:
                with st.expander("📝 Stage 1: Tailored Content"):
                    st.text_area("Stage 1 Output", stage1_result, height=200, key="stage1_display")
            
            with col2:
                with st.expander("🎨 Stage 2: Tone Matched"):
                    st.text_area("Stage 2 Output", stage2_result, height=200, key="stage2_display")
            
            with col3:
                with st.expander("📄 Stage 3: Optimized"):
                    st.text_area("Stage 3 Output", stage3_result, height=200, key="stage3_display")
            
            # Download button
            st.divider()
            
            # Save doc to bytes
            from io import BytesIO
            doc_bytes = BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"Tailored_Resume_{timestamp}.docx"
            
            st.download_button(
                label="📥 Download DOCX Resume",
                data=doc_bytes.getvalue(),
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True
            )
            
            # Preview sections
            st.divider()
            st.markdown("### 📋 Resume Preview")
            
            if sections['summary']:
                with st.expander("Summary", expanded=True):
                    st.markdown(sections['summary'])
            
            if sections['skills']:
                with st.expander("Skills", expanded=False):
                    for category, items in sections['skills'].items():
                        st.markdown(f"**{category}:** {items}")
            
            if sections['experience']:
                with st.expander(f"Experience ({len(sections['experience'])} jobs)", expanded=False):
                    for job in sections['experience']:
                        st.markdown(f"**{job['company']}** | {job['title']} | {job['duration']}")
                        for resp in job['responsibilities']:
                            st.markdown(f"- {resp}")
                        st.divider()
            
            if sections['education']:
                with st.expander("Education", expanded=False):
                    st.markdown(sections['education'])
            
        except Exception as e:
            st.error(f"❌ Error generating resume: {str(e)}")
            st.exception(e)
        
        finally:
            progress_bar.empty()
            status_text.empty()


def run_stage1_tailoring(resume_text, job_description, api_key, additional_notes=""):
    """Stage 1: Tailor content to job requirements"""
    
    prompt = f"""
Analyze this job description and create a tailored version of the resume that emphasizes relevant experience and skills.

**JOB DESCRIPTION:**
{job_description}

**ORIGINAL RESUME:**
{resume_text}

**ADDITIONAL NOTES:**
{additional_notes if additional_notes else "None"}

**TASK:**
1. Identify key requirements from the job description (technologies, skills, responsibilities)
2. Reorder and emphasize relevant skills and experiences
3. Add specific keywords from the job description naturally
4. Keep ALL original job positions and dates intact - ONLY include actual job positions in EXPERIENCE section
5. Quantify achievements where possible
6. DO NOT include general experience descriptions, technology lists, or job description content in the EXPERIENCE section

**OUTPUT FORMAT:**

SUMMARY:
[3-4 sentence summary emphasizing alignment with the role]

SKILLS:
[List skills organized by category, emphasizing job-relevant technologies]

EXPERIENCE:

[Company Name], [Location]
[Job Title] | [Duration]
- [Bullet point 1 - ONLY job-specific achievements and responsibilities]
- [Bullet point 2 - ONLY job-specific achievements and responsibilities]
- [Continue with all bullets - ONLY job-specific content]

[Repeat for all jobs]

**CRITICAL:** The EXPERIENCE section must ONLY contain job positions with their bullet points. DO NOT include:
- General experience descriptions
- Technology lists (those belong in SKILLS)
- Job description content from the posting
- Any text that is not directly about a specific job position

EDUCATION:
[Education details]

Generate the complete tailored resume now:
"""
    
    return call_gemini(prompt, api_key)


def run_stage2_tone_matching(original_resume, tailored_resume, api_key):
    """Stage 2: Match natural writing tone"""
    
    prompt = f"""
Role: You are an expert resume editor. Your task is to revise a "Tailored Resume" to match the writing style and natural language of an "Original Resume" while retaining all critical achievements, keywords, and metrics.

**ORIGINAL RESUME (Style Guide):**
{original_resume}

**TAILORED RESUME (Content Source):**
{tailored_resume}

**TASK AND RULES:**

1. Analyze Tone: Analyze the Original Resume to understand its specific tone. Note the action verbs (e.g., "Developed," "Managed," "Contributed") and the direct, human-like sentence structure.

2. Preserve Key Content: Retain ALL key information from the Tailored Resume:
   - All technologies and skills
   - Specific metrics and achievements
   - Stronger, specific verbs if they describe unique achievements

3. Rewrite for Tone:
   - Make it sound natural and human-written, NOT AI-generated
   - Replace overly corporate buzzwords ("Pioneered," "Championed," "Spearheaded") with direct verbs from the Original Resume ("Developed," "Built," "Enhanced")
   - Keep "Architected" or "Engineered" only if truly architectural work
   - Ensure smooth, straightforward sentence flow

4. ATS-Friendly:
   - Use standard section headings (Summary, Skills, Experience, Education)
   - Spell out all keywords correctly as plain text
   - Use standard bullet points
   - No special characters, columns, or tables

**OUTPUT FORMAT:**

SUMMARY

[Natural, human-written summary paragraph - 3 sentences max]

SKILLS

Frontend Development: [technologies]
Backend Development: [technologies]
Cloud & DevOps: [technologies]
Databases & Storage: [technologies]
Testing & Automation: [technologies]
Version Control & Collaboration: [technologies]
Security: [technologies]
AI/ML & Data Science: [technologies]

EXPERIENCE

[Company Name], [Location]
[Job Title] | [Duration]
- [Natural, direct bullet point - ONLY job-specific work]
- [Natural, direct bullet point - ONLY job-specific work]

[Repeat for all jobs]

**CRITICAL:** The EXPERIENCE section must ONLY contain actual job positions. Each entry must have:
- Company name and location
- Job title and duration
- Bullet points describing what was accomplished at THAT specific job

DO NOT include:
- General experience summaries
- Technology lists (put those in SKILLS section)
- Content from the job description posting
- Any non-job-specific descriptions

EDUCATION

[Degree]
[University] | [Dates]

**CRITICAL:** Output ONLY the revised resume. Start immediately with "SUMMARY"

Generate the resume now:
"""
    
    result = call_gemini(prompt, api_key)
    
    # Clean up markdown
    result = re.sub(r'```.*?\n', '', result)
    result = re.sub(r'```', '', result)
    
    return result


def run_stage3_optimization(tone_matched_resume, api_key, max_bullets=5):
    """Stage 3: Optimize for 2 pages"""
    
    prompt = f"""
You are an expert hiring manager and resume optimization specialist.

**CRITICAL REQUIREMENTS:**
1. The final resume MUST fit on 2 pages maximum (approximately 50-60 lines total)
2. NO markdown formatting (no asterisks, no bold markers like ** **)
3. Plain text only with clean formatting
4. Keep only the most impactful bullet points ({max_bullets} per job maximum)
5. Summary must be 2-3 sentences ONLY

**YOUR TASK:**

Review this resume and optimize it for 2 pages:

{tone_matched_resume}

**OPTIMIZATION RULES:**

Summary:
- Reduce to 2-3 concise sentences
- Focus on years of experience, key technologies, and top achievement

Skills:
- Keep all categories but make them more concise
- Remove redundant tools
- One line per category

Experience:
- ONLY list actual job positions: company, title, duration, and bullet points
- NO introductory paragraphs or technology summaries before the bullets
- NO general descriptions like "Responsibilities include..." or "Technologies used..."
- NO general experience sections - ONLY specific job positions
- NO technology lists in the experience section (those belong in SKILLS)
- NO content from the job description posting itself
- Start directly with company name
- Keep only {max_bullets} MOST IMPACTFUL bullets per job
- Each bullet must start with an action verb and describe what was accomplished at THAT specific job
- Prioritize bullets with metrics and job-relevant technologies
- Each bullet should be 1-2 lines maximum
- The EXPERIENCE section should ONLY contain job positions from the original resume, tailored to the job description

Education:
- Keep minimal - just degree, university, and dates

**OUTPUT FORMAT (NO MARKDOWN, PLAIN TEXT ONLY):**

SUMMARY

[2-3 sentence summary]

SKILLS

Frontend Development: [concise list]
Backend Development: [concise list]
Cloud & DevOps: [concise list]
Databases & Storage: [concise list]
Testing & Automation: [concise list]
Version Control & Collaboration: [concise list]
Security: [concise list]
AI/ML & Data Science: [concise list]

EXPERIENCE

[Company Name], [Location]
[Job Title] | [Duration]
- [Action verb] [impactful achievement with metric]
- [Action verb] [impactful achievement with metric]
- [Action verb] [impactful achievement with metric]
- [Action verb] [impactful achievement with metric]
- [Action verb] [impactful achievement with metric]

[Company Name], [Location]
[Job Title] | [Duration]
- [Bullet 1]
- [Bullet 2]
- [Bullet 3]
- [Bullet 4]

[Company Name], [Location]
[Job Title] | [Duration]
- [Bullet 1]
- [Bullet 2]
- [Bullet 3]
- [Bullet 4]

EDUCATION

[Degree]
[University] | [Dates]

[Degree]
[University] | [Dates]

**CRITICAL RULES:** 
- NO asterisks or markdown formatting
- NO introductory text in Experience section - go straight to bullet points
- NO paragraphs describing the role - ONLY bullet points
- EXPERIENCE section must ONLY contain actual job positions (company, title, duration, bullets)
- DO NOT include general experience descriptions, technology lists, or job description content
- Plain text only
- Must fit on 2 pages (50-60 total lines)
- Start immediately with "SUMMARY"

Generate the optimized 2-page resume now:
"""
    
    result = call_gemini(prompt, api_key)
    
    # Aggressive markdown cleanup
    result = re.sub(r'\*\*([^*]+)\*\*', r'\1', result)
    result = re.sub(r'\*([^*]+)\*', r'\1', result)
    result = re.sub(r'```.*?\n', '', result)
    result = re.sub(r'```', '', result)
    result = re.sub(r'#{1,6}\s', '', result)
    
    return result


def parse_resume_content(content):
    """Parse resume content into structured sections"""
    
    sections = {
        'summary': '',
        'skills': {},
        'experience': [],
        'education': ''
    }
    
    content = content.strip()
    content = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)
    content = re.sub(r'\*([^*]+)\*', r'\1', content)
    
    # Extract sections using regex
    summary_match = re.search(r'SUMMARY:?\s*(.*?)(?=\n\s*SKILLS:?|$)', content, re.DOTALL | re.IGNORECASE)
    skills_match = re.search(r'SKILLS:?\s*(.*?)(?=\n\s*EXPERIENCE:?|$)', content, re.DOTALL | re.IGNORECASE)
    experience_match = re.search(r'EXPERIENCE:?\s*(.*?)(?=\n\s*EDUCATION:?|$)', content, re.DOTALL | re.IGNORECASE)
    education_match = re.search(r'EDUCATION:?\s*(.*?)$', content, re.DOTALL | re.IGNORECASE)
    
    # Extract Summary
    if summary_match:
        sections['summary'] = summary_match.group(1).strip()
    
    # Extract Skills
    if skills_match:
        skills_text = skills_match.group(1).strip()
        for line in skills_text.split('\n'):
            line = line.strip()
            if ':' in line and line:
                parts = line.split(':', 1)
                category = parts[0].strip()
                items = parts[1].strip()
                sections['skills'][category] = items
    
    # Extract Experience
    if experience_match:
        exp_text = experience_match.group(1).strip()
        lines = exp_text.split('\n')
        current_job = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Skip non-job content
            if len(line) > 150 and not any(marker in line for marker in ['-', '•', ',', '|']):
                continue
            
            # Company line (has comma, no bullet, no pipe)
            if ',' in line and not line.startswith(('-', '•')) and '|' not in line:
                if len(line) > 100:
                    continue
                
                if current_job and current_job.get('company') and current_job.get('title'):
                    sections['experience'].append(current_job)
                
                current_job = {
                    'company': line,
                    'title': '',
                    'duration': '',
                    'responsibilities': []
                }
            
            # Title line (has pipe separator)
            elif current_job and '|' in line and not line.startswith(('-', '•')):
                parts = line.split('|')
                if len(parts) >= 2:
                    current_job['title'] = parts[0].strip()
                    current_job['duration'] = parts[1].strip()
            
            # Bullet point
            elif current_job and line.startswith(('-', '•')):
                bullet = line.lstrip('-•').strip()
                if bullet and len(bullet) > 10:
                    current_job['responsibilities'].append(bullet)
        
        if current_job and current_job.get('company') and current_job.get('title'):
            sections['experience'].append(current_job)
    
    # Extract Education
    if education_match:
        sections['education'] = education_match.group(1).strip()
    
    return sections


def create_docx_resume_in_memory(sections, name, phone, email):
    """Create formatted DOCX resume in memory"""
    
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Header - Name
    name_para = doc.add_paragraph(name.upper())
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.runs[0]
    name_run.font.size = Pt(14)
    name_run.font.bold = True
    name_run.font.name = 'Calibri'
    name_para.space_after = Pt(2)
    
    # Contact Info
    contact_info = f'Phone: {phone} | Email: {email}'
    contact_para = doc.add_paragraph(contact_info)
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_para.runs[0]
    contact_run.font.size = Pt(11)
    contact_run.font.name = 'Calibri'
    contact_para.space_after = Pt(6)
    
    # Summary Section
    if sections['summary']:
        summary_heading = doc.add_paragraph('SUMMARY')
        summary_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        summary_heading_run = summary_heading.runs[0]
        summary_heading_run.font.size = Pt(12)
        summary_heading_run.font.bold = True
        summary_heading_run.font.name = 'Calibri'
        summary_heading.space_before = Pt(6)
        summary_heading.space_after = Pt(3)
        
        summary_para = doc.add_paragraph(sections['summary'])
        summary_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        summary_run = summary_para.runs[0]
        summary_run.font.size = Pt(11)
        summary_run.font.name = 'Calibri'
        summary_para.space_after = Pt(6)
    
    # Skills Section
    if sections['skills']:
        skills_heading = doc.add_paragraph('SKILLS')
        skills_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        skills_heading_run = skills_heading.runs[0]
        skills_heading_run.font.size = Pt(12)
        skills_heading_run.font.bold = True
        skills_heading_run.font.name = 'Calibri'
        skills_heading.space_before = Pt(6)
        skills_heading.space_after = Pt(3)
        
        # Two-column table
        skills_list = list(sections['skills'].items())
        num_rows = (len(skills_list) + 1) // 2
        
        skills_table = doc.add_table(rows=num_rows, cols=2)
        
        # Remove borders
        tbl = skills_table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = CT_TblPr.new()
            tbl.tblPr = tblPr
        
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        # Populate skills
        for idx, (category, items) in enumerate(skills_list):
            row = idx // 2
            col = idx % 2
            
            if row >= len(skills_table.rows):
                continue
            
            cell = skills_table.rows[row].cells[col]
            cell.paragraphs[0].clear()
            
            para = cell.paragraphs[0]
            category_run = para.add_run(f"{category}: ")
            category_run.font.bold = True
            category_run.font.size = Pt(11)
            category_run.font.name = 'Calibri'
            
            items_run = para.add_run(items)
            items_run.font.size = Pt(11)
            items_run.font.name = 'Calibri'
            para.space_after = Pt(1)
        
        doc.add_paragraph().space_after = Pt(6)
    
    # Experience Section
    if sections['experience']:
        exp_heading = doc.add_paragraph('EXPERIENCE')
        exp_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        exp_heading_run = exp_heading.runs[0]
        exp_heading_run.font.size = Pt(12)
        exp_heading_run.font.bold = True
        exp_heading_run.font.name = 'Calibri'
        exp_heading.space_before = Pt(6)
        exp_heading.space_after = Pt(3)
        
        for idx, job in enumerate(sections['experience']):
            company_location = job['company']
            title = job.get('title', '').strip()
            duration = job.get('duration', '').strip()
            
            if title and duration:
                job_header = f"{company_location} | {title} | {duration}"
            elif title:
                job_header = f"{company_location} | {title}"
            else:
                job_header = company_location
            
            job_para = doc.add_paragraph(job_header)
            job_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            job_run = job_para.runs[0]
            job_run.font.size = Pt(11)
            job_run.font.name = 'Calibri'
            job_run.font.bold = True
            job_para.space_after = Pt(2)
            
            for resp in job['responsibilities']:
                if resp.strip():
                    bullet_para = doc.add_paragraph(resp, style='List Bullet')
                    bullet_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    bullet_para.paragraph_format.left_indent = Inches(0.25)
                    bullet_para.paragraph_format.first_line_indent = Inches(-0.25)
                    
                    bullet_run = bullet_para.runs[0]
                    bullet_run.font.size = Pt(11)
                    bullet_run.font.name = 'Calibri'
                    bullet_para.space_after = Pt(1)
            
            if idx < len(sections['experience']) - 1:
                spacer = doc.add_paragraph()
                spacer.space_after = Pt(4)
    
    # Education Section
    if sections['education']:
        edu_heading = doc.add_paragraph('EDUCATION')
        edu_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        edu_heading_run = edu_heading.runs[0]
        edu_heading_run.font.size = Pt(12)
        edu_heading_run.font.bold = True
        edu_heading_run.font.name = 'Calibri'
        edu_heading.space_before = Pt(6)
        edu_heading.space_after = Pt(3)
        
        edu_lines = sections['education'].split('\n')
        for line in edu_lines:
            line = line.strip()
            if line:
                edu_para = doc.add_paragraph(line)
                edu_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                edu_run = edu_para.runs[0]
                edu_run.font.size = Pt(11)
                edu_run.font.name = 'Calibri'
                edu_para.space_after = Pt(1)
    
    return doc